from __future__ import annotations

import hashlib
import io
import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace

from app.services.resume_parser_service import (
    ResumeParserError,
    parse_resume_bytes,
    parse_resume_path,
    parse_resume_text,
    parse_resume_upload,
)


SAMPLE_RESUME_TEXT = """Rahul Sharma
Guwahati, Assam
rahul.sharma@example.com | +91 98765 43210
https://www.linkedin.com/in/rahul-sharma

PROFESSIONAL SUMMARY
Accounts professional with 6 years of experience in GST, TDS, payroll, bank reconciliation and financial reporting.

SKILLS
Tally, Microsoft Excel, GST, TDS, Payroll, Accounting, Bank Reconciliation

WORK EXPERIENCE
Accounts Manager | ABC Private Limited
January 2022 - Present
Managed GST, payroll and monthly financial reports.

Senior Accountant | XYZ Services
June 2018 - December 2021
Handled bookkeeping, TDS and reconciliation.

EDUCATION
Master of Commerce, Gauhati University, 2018
Bachelor of Commerce, Gauhati University, 2016

CERTIFICATIONS
Tally Prime Certification

LANGUAGES
English, Assamese, Hindi

Notice Period: 30 days
Current CTC: INR 6,00,000 per annum
Expected CTC: INR 7,50,000 per annum
"""


def _minimal_pdf_bytes(text: str) -> bytes:
    """Create a small one-page PDF without adding another test dependency."""

    escaped = (
        text.replace("\\", "\\\\")
        .replace("(", "\\(")
        .replace(")", "\\)")
        .replace("\r", " ")
        .replace("\n", " ")
    )
    content = f"BT /F1 11 Tf 50 760 Td ({escaped}) Tj ET".encode("latin-1")

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(content)).encode("ascii") + b" >>\nstream\n"
        + content
        + b"\nendstream",
    ]

    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for number, obj in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{number} 0 obj\n".encode("ascii"))
        output.extend(obj)
        output.extend(b"\nendobj\n")

    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))

    output.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(output)


class ResumeTextParsingTests(unittest.TestCase):
    def test_extracts_reviewable_candidate_fields(self):
        result = parse_resume_text(SAMPLE_RESUME_TEXT)
        fields = result["fields"]

        self.assertEqual(fields["full_name"], "Rahul Sharma")
        self.assertEqual(fields["email"], "rahul.sharma@example.com")
        self.assertEqual(fields["phone"], "+919876543210")
        self.assertEqual(fields["location"], "Guwahati, Assam")
        self.assertEqual(fields["current_designation"], "Accounts Manager")
        self.assertEqual(fields["current_employer"], "ABC Private Limited")
        self.assertEqual(fields["total_experience_years"], 6.0)
        self.assertEqual(fields["notice_period"], "30 days")
        self.assertEqual(
            fields["links"]["linkedin"],
            "https://www.linkedin.com/in/rahul-sharma",
        )
        self.assertIn("GST", fields["skills"])
        self.assertIn("Payroll", fields["skills"])
        self.assertGreaterEqual(len(fields["education"]), 2)
        self.assertGreaterEqual(len(fields["employment_history"]), 2)
        self.assertTrue(result["requires_manual_review"])
        self.assertEqual(result["confidence"]["email"], "high")

    def test_rejects_text_without_enough_readable_content(self):
        with self.assertRaises(ResumeParserError) as context:
            parse_resume_text("Name only")

        self.assertEqual(context.exception.code, "resume_text_not_found")
        self.assertEqual(context.exception.status_code, 422)


class ResumeByteValidationTests(unittest.TestCase):
    def test_parses_txt_and_returns_source_metadata(self):
        data = SAMPLE_RESUME_TEXT.encode("utf-8")

        result = parse_resume_bytes(
            data,
            filename="Rahul Sharma Resume.txt",
            content_type="text/plain",
        )

        self.assertTrue(result["ok"])
        self.assertEqual(result["source"]["filename"], "Rahul_Sharma_Resume.txt")
        self.assertEqual(result["source"]["extension"], "txt")
        self.assertEqual(result["source"]["size_bytes"], len(data))
        self.assertEqual(result["source"]["sha256"], hashlib.sha256(data).hexdigest())
        self.assertEqual(result["fields"]["email"], "rahul.sharma@example.com")

    def test_rejects_non_binary_data(self):
        with self.assertRaises(ResumeParserError) as context:
            parse_resume_bytes(  # type: ignore[arg-type]
                "not bytes",
                filename="resume.txt",
                content_type="text/plain",
            )

        self.assertEqual(context.exception.code, "invalid_resume_data")
        self.assertEqual(context.exception.status_code, 400)

    def test_rejects_empty_file(self):
        with self.assertRaises(ResumeParserError) as context:
            parse_resume_bytes(
                b"",
                filename="resume.txt",
                content_type="text/plain",
            )

        self.assertEqual(context.exception.code, "empty_resume_file")

    def test_rejects_file_over_configured_size_limit(self):
        with self.assertRaises(ResumeParserError) as context:
            parse_resume_bytes(
                b"a" * 101,
                filename="resume.txt",
                content_type="text/plain",
                max_bytes=100,
            )

        self.assertEqual(context.exception.code, "resume_too_large")
        self.assertEqual(context.exception.status_code, 413)

    def test_rejects_unsupported_extension(self):
        with self.assertRaises(ResumeParserError) as context:
            parse_resume_bytes(
                SAMPLE_RESUME_TEXT.encode("utf-8"),
                filename="resume.exe",
                content_type="application/octet-stream",
            )

        self.assertEqual(context.exception.code, "unsupported_resume_format")
        self.assertEqual(context.exception.status_code, 415)

    def test_rejects_unsupported_content_type(self):
        with self.assertRaises(ResumeParserError) as context:
            parse_resume_bytes(
                SAMPLE_RESUME_TEXT.encode("utf-8"),
                filename="resume.txt",
                content_type="image/png",
            )

        self.assertEqual(context.exception.code, "unsupported_resume_content_type")

    def test_rejects_fake_pdf_signature(self):
        with self.assertRaises(ResumeParserError) as context:
            parse_resume_bytes(
                SAMPLE_RESUME_TEXT.encode("utf-8"),
                filename="resume.pdf",
                content_type="application/pdf",
            )

        self.assertEqual(context.exception.code, "invalid_pdf_signature")
        self.assertEqual(context.exception.status_code, 415)

    def test_rejects_fake_docx_signature(self):
        with self.assertRaises(ResumeParserError) as context:
            parse_resume_bytes(
                SAMPLE_RESUME_TEXT.encode("utf-8"),
                filename="resume.docx",
                content_type=(
                    "application/vnd.openxmlformats-officedocument."
                    "wordprocessingml.document"
                ),
            )

        self.assertEqual(context.exception.code, "invalid_docx_signature")


class ResumeUploadTests(unittest.TestCase):
    def test_upload_stream_position_is_restored_after_parsing(self):
        data = SAMPLE_RESUME_TEXT.encode("utf-8")
        stream = io.BytesIO(data)
        stream.seek(7)
        upload = SimpleNamespace(
            filename="candidate.txt",
            mimetype="text/plain",
            stream=stream,
        )

        result = parse_resume_upload(upload)

        self.assertEqual(stream.tell(), 7)
        self.assertEqual(result["fields"]["email"], "rahul.sharma@example.com")

    def test_upload_requires_a_file(self):
        with self.assertRaises(ResumeParserError) as context:
            parse_resume_upload(None)

        self.assertEqual(context.exception.code, "resume_file_required")

    def test_upload_requires_a_readable_stream(self):
        upload = SimpleNamespace(filename="candidate.txt", mimetype="text/plain")

        with self.assertRaises(ResumeParserError) as context:
            parse_resume_upload(upload)

        self.assertEqual(context.exception.code, "invalid_resume_file")


class ResumeDocumentParsingTests(unittest.TestCase):
    def test_parses_machine_readable_pdf(self):
        try:
            import pypdf  # noqa: F401
        except ImportError:
            self.skipTest("pypdf is not installed")

        pdf_text = (
            "Rahul Sharma rahul.sharma@example.com +91 98765 43210 "
            "Accounts Manager with 6 years of experience in accounting GST payroll"
        )
        result = parse_resume_bytes(
            _minimal_pdf_bytes(pdf_text),
            filename="candidate.pdf",
            content_type="application/pdf",
        )

        self.assertTrue(result["ok"])
        self.assertEqual(result["source"]["extension"], "pdf")
        self.assertEqual(result["fields"]["email"], "rahul.sharma@example.com")
        self.assertEqual(result["fields"]["phone"], "+919876543210")

    def test_parses_docx_paragraphs_and_table_content(self):
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx is not installed")

        document = Document()
        document.add_heading("Rahul Sharma", level=1)
        document.add_paragraph("rahul.sharma@example.com | +91 98765 43210")
        document.add_heading("Professional Summary", level=2)
        document.add_paragraph(
            "Accounts Manager with 6 years of experience in GST and payroll."
        )
        document.add_heading("Skills", level=2)
        document.add_paragraph("Accounting, GST, Payroll, Microsoft Excel")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Location"
        table.rows[0].cells[1].text = "Guwahati, Assam"

        buffer = io.BytesIO()
        document.save(buffer)

        result = parse_resume_bytes(
            buffer.getvalue(),
            filename="candidate.docx",
            content_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )

        self.assertTrue(result["ok"])
        self.assertEqual(result["source"]["extension"], "docx")
        self.assertEqual(result["fields"]["email"], "rahul.sharma@example.com")
        self.assertIn("GST", result["fields"]["skills"])


class ResumePathTests(unittest.TestCase):
    def test_parses_resume_from_existing_path(self):
        with tempfile.TemporaryDirectory() as temporary_directory:
            path = Path(temporary_directory) / "candidate.txt"
            path.write_text(SAMPLE_RESUME_TEXT, encoding="utf-8")

            result = parse_resume_path(path, content_type="text/plain")

        self.assertEqual(result["fields"]["full_name"], "Rahul Sharma")
        self.assertEqual(result["source"]["filename"], "candidate.txt")

    def test_missing_path_returns_controlled_error(self):
        with self.assertRaises(ResumeParserError) as context:
            parse_resume_path("missing-resume-file.txt")

        self.assertEqual(context.exception.code, "resume_file_not_found")
        self.assertEqual(context.exception.status_code, 404)


if __name__ == "__main__":
    unittest.main()