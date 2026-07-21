"""
Local resume parsing service for the YourComate Recruitment module.

Supported formats:
- PDF with machine-readable text
- DOCX
- TXT

This service deliberately uses deterministic extraction instead of sending candidate
resumes to an external AI provider. Parsed fields are suggestions and must be reviewed
by HR before a candidate record is created or updated.
"""

from __future__ import annotations

import hashlib
import io
import os
import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, BinaryIO, Iterable

from werkzeug.utils import secure_filename


PARSER_VERSION = "1.0.0"

ALLOWED_RESUME_EXTENSIONS = {"pdf", "docx", "txt"}
ALLOWED_CONTENT_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain",
    "application/msword",
    "application/zip",
    "application/x-zip-compressed",
    "application/octet-stream",
    "binary/octet-stream",
    "",
}

MAX_RESUME_BYTES = 8 * 1024 * 1024
MAX_PDF_PAGES = 60
MAX_EXTRACTED_TEXT_CHARS = 250_000
MAX_DOCX_ARCHIVE_FILES = 2_000
MAX_DOCX_UNCOMPRESSED_BYTES = 60 * 1024 * 1024
MAX_DOCX_COMPRESSION_RATIO = 250

EMAIL_PATTERN = re.compile(
    r"(?<![\w.+-])([A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,})(?![\w.-])",
    re.IGNORECASE,
)
URL_PATTERN = re.compile(
    r"\b(?:https?://|www\.)[^\s<>()\[\]{}]+",
    re.IGNORECASE,
)
LINKEDIN_PATTERN = re.compile(
    r"\b(?:https?://)?(?:www\.)?linkedin\.com/in/[A-Za-z0-9_%./-]+",
    re.IGNORECASE,
)
GITHUB_PATTERN = re.compile(
    r"\b(?:https?://)?(?:www\.)?github\.com/[A-Za-z0-9_.-]+",
    re.IGNORECASE,
)
INDIAN_PHONE_PATTERN = re.compile(
    r"(?<!\d)(?:\+?91[\s().-]*)?([6-9](?:[\s().-]*\d){9})(?!\d)"
)
INTERNATIONAL_PHONE_PATTERN = re.compile(
    r"(?<!\d)(\+\d{1,3}[\s().-]*(?:\d[\s().-]*){7,13}\d)(?!\d)"
)
YEAR_PATTERN = re.compile(r"\b(?:19|20)\d{2}\b")

SECTION_ALIASES = {
    "summary": {
        "summary",
        "professional summary",
        "profile",
        "profile summary",
        "career summary",
        "professional profile",
        "objective",
        "career objective",
        "about me",
    },
    "skills": {
        "skills",
        "key skills",
        "technical skills",
        "core competencies",
        "competencies",
        "expertise",
        "areas of expertise",
        "professional skills",
        "tools and technologies",
        "technologies",
    },
    "experience": {
        "experience",
        "work experience",
        "professional experience",
        "employment history",
        "work history",
        "career history",
        "professional background",
    },
    "education": {
        "education",
        "educational qualification",
        "educational qualifications",
        "academic qualification",
        "academic qualifications",
        "academic background",
        "academics",
    },
    "certifications": {
        "certification",
        "certifications",
        "licenses and certifications",
        "professional certifications",
        "courses and certifications",
        "training and certifications",
    },
    "projects": {
        "project",
        "projects",
        "key projects",
        "academic projects",
        "professional projects",
    },
    "languages": {
        "language",
        "languages",
        "language proficiency",
        "languages known",
    },
    "achievements": {
        "achievement",
        "achievements",
        "awards",
        "awards and achievements",
        "accomplishments",
    },
}

HEADING_LOOKUP = {
    alias: canonical
    for canonical, aliases in SECTION_ALIASES.items()
    for alias in aliases
}

NAME_BLOCKLIST = {
    "resume",
    "curriculum vitae",
    "cv",
    "profile",
    "professional profile",
    "professional summary",
    "summary",
    "contact",
    "contact details",
    "personal details",
    "career objective",
    "objective",
    "candidate profile",
}

DEGREE_KEYWORDS = {
    "bachelor",
    "master",
    "b.tech",
    "btech",
    "m.tech",
    "mtech",
    "b.e",
    "m.e",
    "b.sc",
    "bsc",
    "m.sc",
    "msc",
    "b.com",
    "bcom",
    "m.com",
    "mcom",
    "bba",
    "mba",
    "bca",
    "mca",
    "phd",
    "doctorate",
    "diploma",
    "higher secondary",
    "secondary",
    "class x",
    "class xii",
    "10th",
    "12th",
}

COMMON_SKILLS = (
    "accounting",
    "accounts payable",
    "accounts receivable",
    "audit",
    "bank reconciliation",
    "bookkeeping",
    "budgeting",
    "business analysis",
    "cash flow management",
    "cost accounting",
    "financial analysis",
    "financial reporting",
    "gst",
    "income tax",
    "payroll",
    "tally",
    "tds",
    "taxation",
    "microsoft excel",
    "advanced excel",
    "microsoft office",
    "power bi",
    "tableau",
    "data analysis",
    "data visualization",
    "sql",
    "mysql",
    "postgresql",
    "mongodb",
    "python",
    "java",
    "javascript",
    "typescript",
    "react",
    "angular",
    "vue",
    "node.js",
    "flask",
    "django",
    "fastapi",
    "php",
    "laravel",
    "html",
    "css",
    "git",
    "github",
    "docker",
    "kubernetes",
    "aws",
    "azure",
    "google cloud",
    "machine learning",
    "deep learning",
    "artificial intelligence",
    "natural language processing",
    "computer vision",
    "project management",
    "product management",
    "operations management",
    "vendor management",
    "inventory management",
    "supply chain management",
    "procurement",
    "quality assurance",
    "quality control",
    "customer service",
    "customer relationship management",
    "sales",
    "business development",
    "digital marketing",
    "social media marketing",
    "content writing",
    "seo",
    "recruitment",
    "talent acquisition",
    "employee engagement",
    "performance management",
    "training and development",
    "human resources",
    "communication",
    "leadership",
    "team management",
    "problem solving",
    "time management",
    "negotiation",
    "presentation",
)

MONTH_LOOKUP = {
    "jan": 1,
    "january": 1,
    "feb": 2,
    "february": 2,
    "mar": 3,
    "march": 3,
    "apr": 4,
    "april": 4,
    "may": 5,
    "jun": 6,
    "june": 6,
    "jul": 7,
    "july": 7,
    "aug": 8,
    "august": 8,
    "sep": 9,
    "sept": 9,
    "september": 9,
    "oct": 10,
    "october": 10,
    "nov": 11,
    "november": 11,
    "dec": 12,
    "december": 12,
}

DATE_RANGE_PATTERN = re.compile(
    r"\b(?:(?P<start_month>jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|"
    r"jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|"
    r"nov(?:ember)?|dec(?:ember)?)\s+)?(?P<start_year>(?:19|20)\d{2})\s*"
    r"(?:-|–|—|to)\s*"
    r"(?:(?P<end_month>jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|"
    r"jun(?:e)?|jul(?:y)?|aug(?:ust)?|sep(?:t(?:ember)?)?|oct(?:ober)?|"
    r"nov(?:ember)?|dec(?:ember)?)\s+)?"
    r"(?P<end_year>(?:19|20)\d{2}|present|current|till date|date)\b",
    re.IGNORECASE,
)


class ResumeParserError(RuntimeError):
    """Controlled error returned by the Recruitment API to the frontend."""

    def __init__(
        self,
        message: str,
        *,
        code: str = "resume_parse_error",
        status_code: int = 422,
        details: dict[str, Any] | None = None,
    ) -> None:
        super().__init__(message)
        self.message = message
        self.code = code
        self.status_code = status_code
        self.details = details or {}

    def to_dict(self) -> dict[str, Any]:
        return {
            "ok": False,
            "message": self.message,
            "code": self.code,
            "details": self.details,
        }


def safe_str(value: Any) -> str:
    return str(value or "").strip()


def normalize_content_type(value: Any) -> str:
    return safe_str(value).split(";", 1)[0].strip().lower()


def normalize_filename(filename: Any) -> str:
    cleaned = secure_filename(safe_str(filename))
    return cleaned or "resume"


def resume_extension(filename: Any) -> str:
    name = safe_str(filename)
    if "." not in name:
        return ""
    return name.rsplit(".", 1)[-1].lower().strip()


def _read_limited_stream(stream: BinaryIO, max_bytes: int) -> bytes:
    max_bytes = max(int(max_bytes or MAX_RESUME_BYTES), 1)
    original_position = None

    try:
        original_position = stream.tell()
    except (AttributeError, OSError):
        original_position = None

    try:
        data = stream.read(max_bytes + 1)
    except Exception as exc:
        raise ResumeParserError(
            "The uploaded resume could not be read.",
            code="resume_read_failed",
            status_code=400,
            details={"reason": safe_str(exc)},
        ) from exc
    finally:
        if original_position is not None:
            try:
                stream.seek(original_position)
            except (AttributeError, OSError):
                pass

    if not isinstance(data, (bytes, bytearray)):
        raise ResumeParserError(
            "The uploaded resume is not a valid binary file.",
            code="invalid_resume_file",
            status_code=400,
        )

    data = bytes(data)

    if len(data) > max_bytes:
        raise ResumeParserError(
            f"The resume is too large. Maximum allowed size is {max_bytes // (1024 * 1024)} MB.",
            code="resume_too_large",
            status_code=413,
            details={"max_bytes": max_bytes},
        )

    if not data:
        raise ResumeParserError(
            "The uploaded resume is empty.",
            code="empty_resume_file",
            status_code=400,
        )

    return data


def _validate_extension_and_content_type(
    filename: str,
    content_type: str,
) -> str:
    extension = resume_extension(filename)

    if extension not in ALLOWED_RESUME_EXTENSIONS:
        raise ResumeParserError(
            "Unsupported resume format. Upload a PDF, DOCX or TXT file.",
            code="unsupported_resume_format",
            status_code=415,
            details={
                "allowed_extensions": sorted(ALLOWED_RESUME_EXTENSIONS),
                "received_extension": extension,
            },
        )

    normalized_content_type = normalize_content_type(content_type)
    if normalized_content_type not in ALLOWED_CONTENT_TYPES:
        raise ResumeParserError(
            "The uploaded file type does not match an accepted resume format.",
            code="unsupported_resume_content_type",
            status_code=415,
            details={"received_content_type": normalized_content_type},
        )

    return extension


def _validate_file_signature(data: bytes, extension: str) -> None:
    if extension == "pdf" and not data.lstrip().startswith(b"%PDF-"):
        raise ResumeParserError(
            "The file extension is PDF, but the file content is not a valid PDF.",
            code="invalid_pdf_signature",
            status_code=415,
        )

    if extension == "docx" and not data.startswith((b"PK\x03\x04", b"PK\x05\x06", b"PK\x07\x08")):
        raise ResumeParserError(
            "The file extension is DOCX, but the file content is not a valid Word document.",
            code="invalid_docx_signature",
            status_code=415,
        )


def _validate_docx_archive(data: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as archive:
            members = archive.infolist()
            member_names = {member.filename for member in members}

            required_members = {"[Content_Types].xml", "word/document.xml"}
            if not required_members.issubset(member_names):
                raise ResumeParserError(
                    "The uploaded DOCX file is incomplete or damaged.",
                    code="invalid_docx_archive",
                    status_code=422,
                )

            if len(members) > MAX_DOCX_ARCHIVE_FILES:
                raise ResumeParserError(
                    "The DOCX file contains too many internal files.",
                    code="unsafe_docx_archive",
                    status_code=422,
                )

            total_uncompressed = sum(max(member.file_size, 0) for member in members)
            if total_uncompressed > MAX_DOCX_UNCOMPRESSED_BYTES:
                raise ResumeParserError(
                    "The DOCX file expands beyond the permitted processing limit.",
                    code="unsafe_docx_archive",
                    status_code=422,
                )

            for member in members:
                if member.flag_bits & 0x1:
                    raise ResumeParserError(
                        "Password-protected DOCX resumes are not supported.",
                        code="encrypted_docx_not_supported",
                        status_code=422,
                    )

                if member.compress_size <= 0 or member.file_size < 1_000_000:
                    continue

                ratio = member.file_size / member.compress_size
                if ratio > MAX_DOCX_COMPRESSION_RATIO:
                    raise ResumeParserError(
                        "The DOCX file has an unsafe compression ratio.",
                        code="unsafe_docx_archive",
                        status_code=422,
                    )
    except ResumeParserError:
        raise
    except (zipfile.BadZipFile, OSError) as exc:
        raise ResumeParserError(
            "The uploaded DOCX file is damaged or unreadable.",
            code="invalid_docx_archive",
            status_code=422,
            details={"reason": safe_str(exc)},
        ) from exc


def _clean_line(value: Any) -> str:
    line = safe_str(value).replace("\u00a0", " ")
    line = re.sub(r"[\t\r\f\v]+", " ", line)
    line = re.sub(r" {2,}", " ", line)
    line = line.strip(" \t•●▪◦■□◆◇►▸➢➤")
    return line.strip()


def _normalize_extracted_text(text: Any) -> str:
    raw_text = str(text or "").replace("\x00", "")
    raw_lines = raw_text.splitlines()
    output: list[str] = []
    previous_blank = False

    for raw_line in raw_lines:
        line = _clean_line(raw_line)

        if not line:
            if output and not previous_blank:
                output.append("")
            previous_blank = True
            continue

        output.append(line)
        previous_blank = False

    while output and not output[-1]:
        output.pop()

    normalized = "\n".join(output).strip()
    return normalized[:MAX_EXTRACTED_TEXT_CHARS]


def _dedupe(values: Iterable[Any], *, case_insensitive: bool = True) -> list[str]:
    result: list[str] = []
    seen: set[str] = set()

    for value in values:
        text = _clean_line(value)
        if not text:
            continue

        key = text.casefold() if case_insensitive else text
        if key in seen:
            continue

        seen.add(key)
        result.append(text)

    return result


def _extract_pdf_text(data: bytes) -> tuple[str, dict[str, Any], list[str]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ResumeParserError(
            "PDF parsing is not installed on the server.",
            code="missing_pdf_parser_dependency",
            status_code=500,
            details={"required_package": "pypdf"},
        ) from exc

    warnings: list[str] = []

    try:
        reader = PdfReader(io.BytesIO(data), strict=False)
    except Exception as exc:
        raise ResumeParserError(
            "The PDF resume is damaged or unreadable.",
            code="invalid_pdf_file",
            status_code=422,
            details={"reason": safe_str(exc)},
        ) from exc

    if getattr(reader, "is_encrypted", False):
        try:
            decrypt_result = reader.decrypt("")
        except Exception as exc:
            raise ResumeParserError(
                "Password-protected PDF resumes are not supported.",
                code="encrypted_pdf_not_supported",
                status_code=422,
            ) from exc

        if not decrypt_result:
            raise ResumeParserError(
                "Password-protected PDF resumes are not supported.",
                code="encrypted_pdf_not_supported",
                status_code=422,
            )

    page_count = len(reader.pages)
    if page_count > MAX_PDF_PAGES:
        raise ResumeParserError(
            f"The PDF has too many pages. Maximum supported pages: {MAX_PDF_PAGES}.",
            code="resume_page_limit_exceeded",
            status_code=422,
            details={"page_count": page_count, "max_pages": MAX_PDF_PAGES},
        )

    page_texts: list[str] = []
    failed_pages: list[int] = []

    for index, page in enumerate(reader.pages, start=1):
        try:
            extracted = page.extract_text() or ""
        except Exception:
            failed_pages.append(index)
            extracted = ""

        if extracted.strip():
            page_texts.append(extracted)

    if failed_pages:
        warnings.append(
            "Text could not be read from PDF page(s): "
            + ", ".join(str(page) for page in failed_pages)
            + "."
        )

    return (
        _normalize_extracted_text("\n\n".join(page_texts)),
        {"page_count": page_count, "failed_pages": failed_pages},
        warnings,
    )


def _docx_table_lines(table: Any) -> list[str]:
    lines: list[str] = []

    for row in getattr(table, "rows", []):
        cells = []
        for cell in getattr(row, "cells", []):
            cell_text = _clean_line(getattr(cell, "text", ""))
            if cell_text:
                cells.append(cell_text)

        row_text = " | ".join(_dedupe(cells))
        if row_text:
            lines.append(row_text)

    return lines


def _extract_docx_text(data: bytes) -> tuple[str, dict[str, Any], list[str]]:
    _validate_docx_archive(data)

    try:
        from docx import Document
    except ImportError as exc:
        raise ResumeParserError(
            "DOCX parsing is not installed on the server.",
            code="missing_docx_parser_dependency",
            status_code=500,
            details={"required_package": "python-docx"},
        ) from exc

    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:
        raise ResumeParserError(
            "The DOCX resume is damaged or unreadable.",
            code="invalid_docx_file",
            status_code=422,
            details={"reason": safe_str(exc)},
        ) from exc

    lines: list[str] = []

    for paragraph in document.paragraphs:
        text = _clean_line(paragraph.text)
        if text:
            lines.append(text)

    for table in document.tables:
        lines.extend(_docx_table_lines(table))

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            text = _clean_line(paragraph.text)
            if text:
                lines.append(text)

        for paragraph in section.footer.paragraphs:
            text = _clean_line(paragraph.text)
            if text:
                lines.append(text)

    return (
        _normalize_extracted_text("\n".join(_dedupe(lines, case_insensitive=False))),
        {
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
            "section_count": len(document.sections),
        },
        [],
    )


def _extract_txt_text(data: bytes) -> tuple[str, dict[str, Any], list[str]]:
    encodings = ("utf-8-sig", "utf-16", "cp1252", "latin-1")
    decoded = None
    used_encoding = None

    for encoding in encodings:
        try:
            candidate = data.decode(encoding)
        except UnicodeDecodeError:
            continue

        if candidate.strip():
            decoded = candidate
            used_encoding = encoding
            break

    if decoded is None:
        raise ResumeParserError(
            "The TXT resume encoding could not be read.",
            code="invalid_txt_encoding",
            status_code=422,
        )

    return (
        _normalize_extracted_text(decoded),
        {"encoding": used_encoding},
        [],
    )


def _canonical_heading(line: str) -> str:
    cleaned = re.sub(r"[^a-z0-9&/ ]+", " ", line.casefold())
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return HEADING_LOOKUP.get(cleaned, "")


def _extract_sections(text: str) -> dict[str, list[str]]:
    sections: dict[str, list[str]] = {key: [] for key in SECTION_ALIASES}
    current_section = ""

    for raw_line in text.splitlines():
        line = _clean_line(raw_line)

        if not line:
            if current_section and sections[current_section] and sections[current_section][-1] != "":
                sections[current_section].append("")
            continue

        heading = _canonical_heading(line.rstrip(":"))
        if heading:
            current_section = heading
            continue

        if current_section:
            sections[current_section].append(line)

    return {
        key: values
        for key, values in sections.items()
        if any(_clean_line(value) for value in values)
    }


def _section_text(sections: dict[str, list[str]], name: str) -> str:
    values = sections.get(name) or []
    return "\n".join(value for value in values if _clean_line(value)).strip()


def _extract_emails(text: str) -> list[str]:
    return _dedupe(match.group(1).lower() for match in EMAIL_PATTERN.finditer(text))


def _normalize_phone(raw_phone: str) -> str:
    raw_phone = safe_str(raw_phone)
    digits = re.sub(r"\D", "", raw_phone)

    if len(digits) == 12 and digits.startswith("91"):
        return "+91" + digits[2:]

    if len(digits) == 10 and digits[0] in "6789":
        return "+91" + digits

    if raw_phone.startswith("+") and 8 <= len(digits) <= 15:
        return "+" + digits

    return digits if 8 <= len(digits) <= 15 else ""


def _extract_phones(text: str) -> list[str]:
    phones: list[str] = []

    for match in INDIAN_PHONE_PATTERN.finditer(text):
        normalized = _normalize_phone(match.group(0))
        if normalized:
            phones.append(normalized)

    for match in INTERNATIONAL_PHONE_PATTERN.finditer(text):
        normalized = _normalize_phone(match.group(1))
        if normalized:
            phones.append(normalized)

    return _dedupe(phones)


def _clean_url(url: str) -> str:
    return safe_str(url).rstrip(".,;:!?)\"]}")


def _extract_links(text: str) -> dict[str, Any]:
    linkedin = _dedupe(_clean_url(match.group(0)) for match in LINKEDIN_PATTERN.finditer(text))
    github = _dedupe(_clean_url(match.group(0)) for match in GITHUB_PATTERN.finditer(text))
    all_urls = _dedupe(_clean_url(match.group(0)) for match in URL_PATTERN.finditer(text))

    other = [
        url
        for url in all_urls
        if not any(url.casefold() == item.casefold() for item in linkedin + github)
    ]

    return {
        "linkedin": linkedin[0] if linkedin else "",
        "github": github[0] if github else "",
        "other": other[:10],
    }


def _looks_like_name(line: str) -> bool:
    candidate = _clean_line(line)
    lowered = candidate.casefold()

    if not candidate or lowered in NAME_BLOCKLIST:
        return False

    if len(candidate) < 3 or len(candidate) > 70:
        return False

    if EMAIL_PATTERN.search(candidate) or URL_PATTERN.search(candidate):
        return False

    if INDIAN_PHONE_PATTERN.search(candidate) or INTERNATIONAL_PHONE_PATTERN.search(candidate):
        return False

    if any(char.isdigit() for char in candidate):
        return False

    words = candidate.split()
    if not 2 <= len(words) <= 6:
        return False

    allowed_words = 0
    for word in words:
        token = word.strip(".,'-")
        if token and all(char.isalpha() or char in "'-" for char in token):
            allowed_words += 1

    if allowed_words != len(words):
        return False

    uppercase_ratio = sum(char.isupper() for char in candidate) / max(
        sum(char.isalpha() for char in candidate),
        1,
    )
    title_case_words = sum(word[:1].isupper() for word in words)

    return uppercase_ratio >= 0.65 or title_case_words >= max(len(words) - 1, 2)


def _extract_full_name(text: str) -> tuple[str, str]:
    lines = [_clean_line(line) for line in text.splitlines() if _clean_line(line)]

    explicit_patterns = (
        re.compile(r"^(?:full\s+name|candidate\s+name|name)\s*[:\-]\s*(.+)$", re.IGNORECASE),
        re.compile(r"^i\s+am\s+([A-Za-z][A-Za-z .'-]{2,65})$", re.IGNORECASE),
    )

    for line in lines[:30]:
        for pattern in explicit_patterns:
            match = pattern.match(line)
            if match and _looks_like_name(match.group(1)):
                return _clean_line(match.group(1)).title(), "high"

    for line in lines[:15]:
        if _looks_like_name(line):
            normalized = line.title() if line.isupper() else line
            return normalized, "medium"

    return "", "low"


def _extract_labeled_value(text: str, labels: Iterable[str], max_length: int = 160) -> str:
    escaped = "|".join(re.escape(label) for label in labels)
    pattern = re.compile(
        rf"^(?:{escaped})\s*[:\-]\s*(.+)$",
        re.IGNORECASE | re.MULTILINE,
    )
    match = pattern.search(text)

    if not match:
        return ""

    return _clean_line(match.group(1))[:max_length]


def _extract_location(text: str) -> str:
    labeled = _extract_labeled_value(
        text,
        ("current location", "location", "city", "address"),
        max_length=220,
    )
    if labeled:
        return labeled

    top_lines = [_clean_line(line) for line in text.splitlines()[:25] if _clean_line(line)]
    for line in top_lines:
        if re.search(r"\b(?:assam|guwahati|delhi|mumbai|kolkata|bengaluru|bangalore|"
                     r"hyderabad|chennai|pune|noida|gurugram|gurgaon|india)\b", line, re.IGNORECASE):
            if len(line) <= 140 and not EMAIL_PATTERN.search(line):
                return line

    return ""


def _split_list_items(values: Iterable[str], *, max_items: int = 40) -> list[str]:
    items: list[str] = []

    for value in values:
        line = _clean_line(value)
        if not line:
            continue

        parts = re.split(r"\s*[|;,•●▪◦]\s*", line)
        if len(parts) == 1 and "," not in line and len(line.split()) > 12:
            continue

        for part in parts:
            item = _clean_line(part)
            if not item or len(item) > 80:
                continue
            if 1 <= len(item.split()) <= 8:
                items.append(item)

    return _dedupe(items)[:max_items]


def _extract_skills(text: str, sections: dict[str, list[str]]) -> list[str]:
    skills = _split_list_items(sections.get("skills") or [], max_items=60)
    lowered_text = text.casefold()

    for skill in COMMON_SKILLS:
        if re.search(rf"(?<![a-z0-9]){re.escape(skill.casefold())}(?![a-z0-9])", lowered_text):
            skills.append(skill)

    return _dedupe(skills)[:60]


def _group_blocks(values: Iterable[str], *, max_blocks: int = 25) -> list[str]:
    blocks: list[str] = []
    current: list[str] = []

    for value in values:
        line = _clean_line(value)

        if not line:
            if current:
                blocks.append(" | ".join(current))
                current = []
            continue

        if current and (YEAR_PATTERN.search(line) or len(current) >= 4):
            blocks.append(" | ".join(current))
            current = [line]
        else:
            current.append(line)

    if current:
        blocks.append(" | ".join(current))

    return _dedupe(blocks)[:max_blocks]


def _extract_education(sections: dict[str, list[str]]) -> list[dict[str, Any]]:
    education_text = _section_text(sections, "education")
    if not education_text:
        return []

    entries: list[dict[str, Any]] = []
    blocks = _group_blocks(sections.get("education") or [], max_blocks=20)

    for block in blocks:
        lowered = block.casefold()
        if not any(keyword in lowered for keyword in DEGREE_KEYWORDS) and not YEAR_PATTERN.search(block):
            continue

        years = _dedupe(YEAR_PATTERN.findall(block))
        entries.append({
            "text": block,
            "years": years,
        })

    return entries[:20]


def _extract_employment_history(sections: dict[str, list[str]]) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []

    for block in _group_blocks(sections.get("experience") or [], max_blocks=30):
        date_ranges = [match.group(0) for match in DATE_RANGE_PATTERN.finditer(block)]
        entries.append({
            "text": block,
            "date_ranges": _dedupe(date_ranges),
        })

    return entries[:30]


def _month_index(month_name: str | None, year_value: str, *, end: bool = False) -> int:
    year = int(year_value)
    month = MONTH_LOOKUP.get(safe_str(month_name).casefold(), 12 if end else 1)
    return year * 12 + month - 1


def _merged_month_count(intervals: list[tuple[int, int]]) -> int:
    if not intervals:
        return 0

    intervals = sorted(intervals)
    merged: list[list[int]] = []

    for start, end in intervals:
        if end < start:
            continue

        if not merged or start > merged[-1][1] + 1:
            merged.append([start, end])
        else:
            merged[-1][1] = max(merged[-1][1], end)

    return sum(end - start + 1 for start, end in merged)


def _extract_total_experience(text: str, experience_text: str) -> tuple[float | None, str]:
    explicit_patterns = (
        re.compile(
            r"(?:total\s+)?(?:professional\s+|work\s+)?experience\s*[:\-]?\s*"
            r"(\d{1,2}(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)",
            re.IGNORECASE,
        ),
        re.compile(
            r"(\d{1,2}(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)\s+(?:of\s+)?experience",
            re.IGNORECASE,
        ),
    )

    for pattern in explicit_patterns:
        match = pattern.search(text)
        if match:
            try:
                value = float(match.group(1))
            except ValueError:
                continue

            if 0 <= value <= 60:
                return value, "explicit"

    if not experience_text:
        return None, "not_found"

    now = datetime.now(timezone.utc)
    current_month_index = now.year * 12 + now.month - 1
    intervals: list[tuple[int, int]] = []

    for match in DATE_RANGE_PATTERN.finditer(experience_text):
        start = _month_index(match.group("start_month"), match.group("start_year"))
        end_year = match.group("end_year").casefold()

        if end_year in {"present", "current", "till date", "date"}:
            end = current_month_index
        else:
            end = _month_index(match.group("end_month"), end_year, end=True)

        if start <= end <= current_month_index + 12:
            intervals.append((start, min(end, current_month_index)))

    total_months = _merged_month_count(intervals)
    if total_months <= 0:
        return None, "not_found"

    return round(total_months / 12, 1), "calculated_from_date_ranges"


def _extract_summary(sections: dict[str, list[str]]) -> str:
    summary = _section_text(sections, "summary")
    if not summary:
        return ""

    return summary[:1_500]


def _extract_certifications(sections: dict[str, list[str]]) -> list[str]:
    return _split_list_items(sections.get("certifications") or [], max_items=30)


def _extract_languages(sections: dict[str, list[str]]) -> list[str]:
    excluded_labels = {
        "notice period",
        "joining availability",
        "available to join",
        "current ctc",
        "current salary",
        "present ctc",
        "present salary",
        "expected ctc",
        "expected salary",
        "salary expectation",
    }
    values = []

    for item in _split_list_items(sections.get("languages") or [], max_items=30):
        label = item.split(":", 1)[0].strip().casefold()
        if label in excluded_labels:
            continue
        if any(token in label for token in ("salary", "ctc", "notice period")):
            continue
        values.append(item)

    return _dedupe(values)[:20]


def _extract_salary_fields(text: str) -> dict[str, str]:
    current_salary = _extract_labeled_value(
        text,
        ("current ctc", "current salary", "present ctc", "present salary"),
    )
    expected_salary = _extract_labeled_value(
        text,
        ("expected ctc", "expected salary", "salary expectation"),
    )

    return {
        "current_salary": current_salary,
        "expected_salary": expected_salary,
    }


def _extract_candidate_fields(
    text: str,
    sections: dict[str, list[str]],
) -> tuple[dict[str, Any], dict[str, str]]:
    emails = _extract_emails(text)
    phones = _extract_phones(text)
    full_name, name_confidence = _extract_full_name(text)
    links = _extract_links(text)
    experience_text = _section_text(sections, "experience")
    total_experience, experience_source = _extract_total_experience(text, experience_text)
    salary_fields = _extract_salary_fields(text)

    current_designation = _extract_labeled_value(
        text,
        ("current designation", "present designation", "designation", "current role"),
    )
    current_employer = _extract_labeled_value(
        text,
        ("current employer", "present employer", "current company", "organization", "organisation"),
    )

    employment_history = _extract_employment_history(sections)
    if employment_history and (not current_designation or not current_employer):
        first_entry = safe_str(employment_history[0].get("text"))
        parts = [_clean_line(part) for part in first_entry.split("|") if _clean_line(part)]

        if len(parts) >= 2:
            if not current_designation and not DATE_RANGE_PATTERN.search(parts[0]):
                current_designation = parts[0][:160]
            if not current_employer and not DATE_RANGE_PATTERN.search(parts[1]):
                current_employer = parts[1][:160]

    notice_period = _extract_labeled_value(
        text,
        ("notice period", "joining availability", "available to join"),
    )

    fields = {
        "full_name": full_name,
        "email": emails[0] if emails else "",
        "alternate_emails": emails[1:5],
        "phone": phones[0] if phones else "",
        "alternate_phones": phones[1:5],
        "location": _extract_location(text),
        "professional_summary": _extract_summary(sections),
        "current_designation": current_designation,
        "current_employer": current_employer,
        "total_experience_years": total_experience,
        "notice_period": notice_period,
        "current_salary": salary_fields["current_salary"],
        "expected_salary": salary_fields["expected_salary"],
        "skills": _extract_skills(text, sections),
        "education": _extract_education(sections),
        "employment_history": employment_history,
        "certifications": _extract_certifications(sections),
        "languages": _extract_languages(sections),
        "links": links,
    }

    confidence = {
        "full_name": name_confidence,
        "email": "high" if fields["email"] else "low",
        "phone": "high" if fields["phone"] else "low",
        "location": "medium" if fields["location"] else "low",
        "skills": "medium" if fields["skills"] else "low",
        "education": "medium" if fields["education"] else "low",
        "employment_history": "medium" if fields["employment_history"] else "low",
        "total_experience_years": (
            "high"
            if experience_source == "explicit"
            else "medium"
            if experience_source == "calculated_from_date_ranges"
            else "low"
        ),
        "current_designation": "high" if current_designation else "low",
        "current_employer": "high" if current_employer else "low",
        "notice_period": "high" if notice_period else "low",
    }

    return fields, confidence


def _section_preview(sections: dict[str, list[str]]) -> dict[str, str]:
    return {
        name: _section_text(sections, name)[:5_000]
        for name in sections
        if _section_text(sections, name)
    }


def parse_resume_text(text: str) -> dict[str, Any]:
    """
    Parse already-extracted resume text into reviewable candidate fields.

    This helper is useful for tests and for a future OCR fallback. It does not
    validate or store an uploaded file.
    """

    normalized_text = _normalize_extracted_text(text)
    if len(re.sub(r"\s+", "", normalized_text)) < 20:
        raise ResumeParserError(
            "The resume does not contain enough readable text to parse.",
            code="resume_text_not_found",
            status_code=422,
        )

    sections = _extract_sections(normalized_text)
    fields, confidence = _extract_candidate_fields(normalized_text, sections)

    return {
        "fields": fields,
        "confidence": confidence,
        "sections": _section_preview(sections),
        "raw_text": normalized_text,
        "text_length": len(normalized_text),
        "requires_manual_review": True,
    }


def parse_resume_bytes(
    data: bytes,
    *,
    filename: str,
    content_type: str = "",
    max_bytes: int = MAX_RESUME_BYTES,
) -> dict[str, Any]:
    """Validate and parse resume bytes without writing the file to disk."""

    if not isinstance(data, (bytes, bytearray)):
        raise ResumeParserError(
            "Resume data must be provided as bytes.",
            code="invalid_resume_data",
            status_code=400,
        )

    data = bytes(data)
    if not data:
        raise ResumeParserError(
            "The uploaded resume is empty.",
            code="empty_resume_file",
            status_code=400,
        )

    if len(data) > max_bytes:
        raise ResumeParserError(
            f"The resume is too large. Maximum allowed size is {max_bytes // (1024 * 1024)} MB.",
            code="resume_too_large",
            status_code=413,
            details={"max_bytes": max_bytes},
        )

    safe_filename = normalize_filename(filename)
    extension = _validate_extension_and_content_type(safe_filename, content_type)
    _validate_file_signature(data, extension)

    if extension == "pdf":
        text, document_meta, warnings = _extract_pdf_text(data)
    elif extension == "docx":
        text, document_meta, warnings = _extract_docx_text(data)
    else:
        text, document_meta, warnings = _extract_txt_text(data)

    if len(re.sub(r"\s+", "", text)) < 20:
        details = {
            "extension": extension,
            "filename": safe_filename,
        }
        if extension == "pdf":
            details["possible_reason"] = "The PDF may be image-only or scanned."

        raise ResumeParserError(
            "No usable text could be extracted from the resume. HR can enter the candidate details manually.",
            code="resume_text_not_found",
            status_code=422,
            details=details,
        )

    parsed = parse_resume_text(text)

    if len(text) >= MAX_EXTRACTED_TEXT_CHARS:
        warnings.append(
            f"Extracted text was limited to {MAX_EXTRACTED_TEXT_CHARS} characters."
        )

    return {
        "ok": True,
        "parser_version": PARSER_VERSION,
        "parsed_at": datetime.now(timezone.utc).isoformat(),
        "source": {
            "filename": safe_filename,
            "original_filename": safe_str(filename),
            "extension": extension,
            "content_type": normalize_content_type(content_type),
            "size_bytes": len(data),
            "sha256": hashlib.sha256(data).hexdigest(),
        },
        "document": document_meta,
        **parsed,
        "warnings": _dedupe(warnings),
    }


def parse_resume_upload(
    uploaded_file: Any,
    *,
    filename: str | None = None,
    content_type: str | None = None,
    max_bytes: int = MAX_RESUME_BYTES,
) -> dict[str, Any]:
    """
    Parse a Flask/Werkzeug FileStorage object or another file-like object.

    The stream position is restored after reading so the Recruitment route can
    safely store the same original upload after parsing.
    """

    if uploaded_file is None:
        raise ResumeParserError(
            "Resume file is required.",
            code="resume_file_required",
            status_code=400,
        )

    resolved_filename = (
        safe_str(filename)
        or safe_str(getattr(uploaded_file, "filename", ""))
        or "resume"
    )
    resolved_content_type = (
        safe_str(content_type)
        or safe_str(getattr(uploaded_file, "mimetype", ""))
        or safe_str(getattr(uploaded_file, "content_type", ""))
    )

    stream = getattr(uploaded_file, "stream", uploaded_file)
    if not hasattr(stream, "read"):
        raise ResumeParserError(
            "The uploaded resume is not a readable file.",
            code="invalid_resume_file",
            status_code=400,
        )

    data = _read_limited_stream(stream, max_bytes)

    return parse_resume_bytes(
        data,
        filename=resolved_filename,
        content_type=resolved_content_type,
        max_bytes=max_bytes,
    )


def parse_resume_path(
    file_path: str | os.PathLike[str],
    *,
    original_filename: str | None = None,
    content_type: str = "",
    max_bytes: int = MAX_RESUME_BYTES,
) -> dict[str, Any]:
    """Parse an existing local file, mainly for tests and controlled imports."""

    path = Path(file_path)
    if not path.is_file():
        raise ResumeParserError(
            "The resume file was not found.",
            code="resume_file_not_found",
            status_code=404,
        )

    try:
        with path.open("rb") as file:
            data = _read_limited_stream(file, max_bytes)
    except ResumeParserError:
        raise
    except OSError as exc:
        raise ResumeParserError(
            "The resume file could not be opened.",
            code="resume_read_failed",
            status_code=422,
            details={"reason": safe_str(exc)},
        ) from exc

    return parse_resume_bytes(
        data,
        filename=original_filename or path.name,
        content_type=content_type,
        max_bytes=max_bytes,
    )


__all__ = [
    "ALLOWED_RESUME_EXTENSIONS",
    "MAX_RESUME_BYTES",
    "PARSER_VERSION",
    "ResumeParserError",
    "parse_resume_bytes",
    "parse_resume_path",
    "parse_resume_text",
    "parse_resume_upload",
]